import os, re, time, pathlib, base64
from io import BytesIO
from typing import List, Dict, Tuple, Optional

import numpy as np
import streamlit as st
from openai import OpenAI
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt
import tiktoken
import faiss

# ===================== Brand (logo + CSS) =====================
# REPLACE THIS PATH with your actual logo file if different
LOGO_PATH = "assets/efficio-logo.png"   # <--- put your Efficio logo here (png or svg)

EFFICIO_CSS = """
<style>
:root{
  --ef-text:#0f172a;
  --ef-muted:#64748b;
  --ef-bg:#ffffff;

  --ef-primary:#091a67;       /* Efficio Main Blue */
  --ef-primary-600:#2b335d;   /* Efficio Logo Blue */
  --ef-accent:#00bc9e;        /* Efficio Elements Blue */

  --ef-border:#E6E9EF;
  --ef-shadow:0 16px 40px rgba(16,24,40,.06), 0 2px 6px rgba(16,24,40,.04);
  --ef-radius:18px;
}

html,body,[data-testid="stAppViewContainer"]{
  background:var(--ef-bg)!important; color:var(--ef-text)!important;
}
h1,h2,h3,h4,h5,h6{color:var(--ef-primary)!important;font-weight:800;}
p,li,ul,ol,span,div{color:var(--ef-text)!important;}
label{color:var(--ef-text)!important}

/* ---------- Streamlit TOP BAR / TOOLBAR THEME ---------- */
header[data-testid="stHeader"]{
  background:#ffffff !important;
  border-bottom:1px solid var(--ef-border) !important;
  box-shadow:0 4px 14px rgba(9,26,103,.06) !important;
}
header[data-testid="stHeader"] *{
  color:var(--ef-primary) !important;
  fill:var(--ef-primary) !important;
  stroke:var(--ef-primary) !important;
}
div[data-testid="stDecoration"]{ background:#ffffff !important; }
header[data-testid="stHeader"] a,
header[data-testid="stHeader"] button{
  color:var(--ef-primary) !important;
}
header[data-testid="stHeader"] svg{
  color:var(--ef-primary) !important;
  fill:var(--ef-primary) !important;
  stroke:var(--ef-primary) !important;
}
header [data-testid="stStatusWidget"] > div{
  filter:none !important;
}

/* ---------- HERO (full-width header card) ---------- */
.ef-hero{
  border:1px solid var(--ef-border);
  border-radius:var(--ef-radius);
  background:radial-gradient(1200px 400px at -10% -20%, #f2f6ff 0%, #ffffff 55%) no-repeat;
  padding:26px 28px;
  box-shadow:var(--ef-shadow);
  display:flex;align-items:center;justify-content:space-between;gap:18px;margin:22px 0;
}
.ef-hero-left{display:flex;align-items:center;gap:16px}
.ef-logo-img{height:36px;display:block}
.ef-fallback{height:36px;width:36px;display:block;border-radius:10px}
.ef-title{margin:0;font-size:28px;line-height:1.2;color:var(--ef-primary)}
.ef-subtle{color:var(--ef-muted);font-size:14px;margin-top:2px}
.ef-pill{background:#f4f6ff;color:var(--ef-primary);border:1px solid var(--ef-border);
  border-radius:999px;padding:6px 12px;font-size:12px}

/* ---------- Inputs stay WHITE ---------- */
.stTextInput div[data-baseweb="input"],
.stTextInput div[data-baseweb="input"] > div,
.stTextInput input,
.stTextArea div[data-baseweb="textarea"],
.stTextArea textarea,
.stSelectbox div[data-baseweb="select"],
.stSelectbox div[data-baseweb="select"] > div,
.stSelectbox [role="combobox"]{
  background:#ffffff !important;
  color:var(--ef-text) !important;
  border:1px solid var(--ef-border) !important;
  border-radius:12px !important;
  box-shadow:none !important;
}
.stTextInput div[data-baseweb="input"]:focus-within,
.stTextArea div[data-baseweb="textarea"]:focus-within,
.stSelectbox div[data-baseweb="select"] > div:focus-within{
  border:1px solid var(--ef-accent) !important;
  box-shadow:0 0 0 3px rgba(0,188,158,.18) !important;
}
::placeholder{color:#94a3b8 !important}

/* ---------- Expander cards ---------- */
div[data-testid="stExpander"]{
  border:1px solid var(--ef-border)!important;border-radius:var(--ef-radius)!important;
  background:#fff !important;box-shadow:var(--ef-shadow);overflow:hidden;
}
div[data-testid="stExpander"] summary,
div[data-testid="stExpander"] summary *{
  color:var(--ef-primary)!important;font-weight:700!important
}

/* ---------- Primary buttons (force WHITE text everywhere) ---------- */
.stButton button, .stDownloadButton button,
.stButton > button, .stDownloadButton > button,
.stButton button[kind="primary"], .stDownloadButton button[kind="primary"]{
  color:#ffffff !important;
  background:linear-gradient(180deg, var(--ef-primary) 0%, #0b1a77 100%) !important;
  border:1px solid var(--ef-primary) !important;
  border-radius:12px !important;
  padding:10px 16px !important; height:44px !important; font-weight:700 !important;
  box-shadow:0 6px 16px rgba(9,26,103,.18) !important;
  text-shadow:0 1px 0 rgba(0,0,0,.18) !important;
  transition:all .15s ease-in-out !important;
}
.stButton button *, .stDownloadButton button *{
  color:#ffffff !important;
}
.stButton button:hover, .stDownloadButton button:hover{
  background:var(--ef-primary-600) !important; border-color:var(--ef-primary-600) !important;
  transform:translateY(-1px);
}
.stButton button:focus, .stDownloadButton button:focus{
  outline:none !important; box-shadow:0 0 0 3px rgba(0,188,158,.28)!important;
}

/* ---------- ALWAYS-VISIBLE "Expand all sections" (checkbox styled as a switch) ---------- */
div[data-testid="stCheckbox"] label{
  display:flex; align-items:center; gap:12px;
}
div[data-testid="stCheckbox"] div[role="checkbox"]{
  width:44px !important; height:24px !important;
  border-radius:999px !important;
  background:#eef2f7 !important;
  border:1px solid var(--ef-border) !important;
  position:relative !important;
  box-shadow:none !important;
}
div[data-testid="stCheckbox"] div[role="checkbox"][aria-checked="true"]{
  background:var(--ef-accent) !important; border-color:var(--ef-accent) !important;
}
div[data-testid="stCheckbox"] div[role="checkbox"]::after{
  content:"";
  position:absolute; top:2px; left:2px;
  width:20px; height:20px;
  border-radius:50%;
  background:#fff;
  box-shadow:0 2px 6px rgba(2,6,23,.12);
  transition:left .15s ease;
}
div[data-testid="stCheckbox"] div[role="checkbox"][aria-checked="true"]::after{
  left:22px;
}

/* ---------- Tables ---------- */
table{border-collapse:separate!important;border-spacing:0!important;width:100%!important}
thead th{background:#f7f9ff!important;color:var(--ef-primary)!important;border-bottom:1px solid var(--ef-border)!important}
tbody td{border-bottom:1px solid #f1f3f7!important}

/* ---------- Footer hint ---------- */
.ef-hint{color:var(--ef-muted);font-size:12px;text-align:center;padding:18px 0}

/* =========================================================
   GLOBAL POPUP / MENU OVERRIDES (BaseWeb/Streamlit portals)
   This catches the Deploy/Rerun/Settings dropdown which is
   rendered at the end of <body>, outside #stApp.
   ========================================================= */
div[data-baseweb="popover"] div[role="menu"],
div[role="menu"]{
  background:#ffffff !important;
  color:var(--ef-text) !important;
  border:1px solid var(--ef-border) !important;
  border-radius:12px !important;
  box-shadow:var(--ef-shadow) !important;
}

/* Menu items (anchor, button, li) */
div[data-baseweb="popover"] [role="menuitem"],
[role="menuitem"],
div[data-baseweb="popover"] [role="menuitem"] *,
[role="menuitem"] *{
  color:var(--ef-text) !important;
}

/* Hover / focus state */
div[data-baseweb="popover"] [role="menuitem"]:hover,
[role="menuitem"]:hover{
  background:#f4f6ff !important;
  color:var(--ef-primary) !important;
}

/* Shortcut badges (e.g., “R”, “C”) */
div[data-baseweb="popover"] kbd,
[role="menu"] kbd{
  background:#eef2f9 !important;
  color:var(--ef-primary) !important;
  border:1px solid var(--ef-border) !important;
  box-shadow:none !important;
}
</style>
"""





def render_topbar():
    """Hero header with Efficio logo; uses LOGO_PATH if present."""
    p = pathlib.Path(LOGO_PATH)
    if p.exists():
        mime = "image/svg+xml" if p.suffix.lower() == ".svg" else "image/png"
        with open(p, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
        logo_html = f'<img class="ef-logo-img" alt="Efficio logo" src="data:{mime};base64,{b64}"/>'
    else:
        logo_html = """
        <svg class="ef-fallback" viewBox="0 0 36 36" aria-hidden="true">
          <rect x="0" y="0" width="36" height="36" rx="10" fill="#091a67"/>
          <path d="M25 10H12v3.2h8.7V16H12v3.2h13V22H9.5V7.8H25V10Z" fill="#fff"/>
        </svg>
        """

    st.markdown(f"""
    <div class="ef-hero">
      <div class="ef-hero-left">
        {logo_html}
        <div>
          <h2 class="ef-title">Efficio Value Chain Builder</h2>
          <div class="ef-subtle">Data-driven industry analysis</div>
        </div>
      </div>
      <span class="ef-pill">v1</span>
    </div>""", unsafe_allow_html=True)

# ---------------- Page setup ----------------
st.set_page_config(page_title="Efficio Value Chain Builder", page_icon="🧭", layout="wide")

st.markdown(EFFICIO_CSS, unsafe_allow_html=True)
render_topbar()

# ---------------- Config ----------------
APP_DIR = pathlib.Path(__file__).parent
KNOW_DIR = APP_DIR / "knowledge"
KNOW_DIR.mkdir(exist_ok=True)

API_KEY = "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".strip()
if not API_KEY:
    st.error("OPENAI_API_KEY not found. Add it in code for dev, or use .streamlit/secrets.toml / env var.")
    st.stop()
client = OpenAI(api_key=API_KEY)

CHAT_MODEL = "gpt-4o-mini"
EMBED_MODEL = "text-embedding-3-small"

# ================= Token budgets (+20%) =================
SECTION_BUDGET = {
    "1.1 Introduction": 1920,
    "1.2 Global Trade": 2160,
    "1.2 Case Study (2 Countries)": 2160,
    "2.1 Country Landscape": 1920,
    "2.2 Segments Overview": 1920,
    "2.3 Trade Analysis": 1680,
    "2.4 Local Capabilities": 1680,
    "3.1 Associated Industries": 2160,
    "3.2 Value Chain Analysis": 2400,
    "3.2 Supplier Case Study": 1920,
    "3.3 Raw Material Analysis": 2400,
    "3.3 Raw Material Supplier Case Study": 1920,
    "4. Consolidated Opportunities": 1440,
}
MAX_CONT_ROUNDS = 2
RETRIEVAL_K = 8

# ---------------- Loaders & FAISS ----------------
enc = tiktoken.get_encoding("cl100k_base")

def load_txt(p): return p.read_text(encoding="utf-8", errors="ignore")
def load_docx(p): return "\n".join([para.text for para in Document(p).paragraphs])
def load_pdf(p): return "\n".join([page.extract_text() or "" for page in PdfReader(p).pages])

def read_any(p: pathlib.Path) -> str:
    s = p.suffix.lower()
    if s in (".txt",".md"): return load_txt(p)
    if s==".docx": return load_docx(p)
    if s==".pdf": return load_pdf(p)
    return ""

def chunk_text(text: str, max_tokens=600, overlap=80) -> List[str]:
    toks = enc.encode(text); chunks=[]; i=0
    while i < len(toks):
        j = min(i+max_tokens, len(toks))
        chunk = enc.decode(toks[i:j]).strip()
        if chunk: chunks.append(chunk)
        if j == len(toks): break
        i = max(0, j - overlap)
    return chunks

def _normalize_rows(x: np.ndarray) -> np.ndarray:
    faiss.normalize_L2(x); return x

@st.cache_resource(show_spinner=False)
def build_index(knowledge_dir: pathlib.Path) -> Tuple[faiss.IndexFlatIP, List[Dict]]:
    docs=[]
    for p in knowledge_dir.glob("**/*"):
        if p.is_file():
            raw = read_any(p)
            if not raw.strip(): continue
            for i,ch in enumerate(chunk_text(raw)):
                docs.append({"id": f"{p.name}#chunk{i+1}", "source": p.name, "text": ch})
    if not docs:
        raise RuntimeError("No knowledge found in ./knowledge")

    texts = [d["text"] for d in docs]
    vecs=[]; bs=64
    for k in range(0,len(texts),bs):
        emb = client.embeddings.create(model=EMBED_MODEL, input=texts[k:k+bs])
        vecs.extend([e.embedding for e in emb.data])
    vecs = np.array(vecs, dtype="float32")

    index = faiss.IndexFlatIP(vecs.shape[1])
    _normalize_rows(vecs); index.add(vecs)
    return index, docs

@st.cache_resource(show_spinner=False)
def load_index_cached(): return build_index(KNOW_DIR)

def retrieve(index, docs, query: str, k=RETRIEVAL_K):
    qv = client.embeddings.create(model=EMBED_MODEL, input=[query]).data[0].embedding
    qv = np.array([qv], dtype="float32"); _normalize_rows(qv)
    scores, idxs = index.search(qv, k)
    res=[]
    for sc,i in zip(scores[0], idxs[0]):
        if i == -1: continue
        hit = docs[i].copy(); hit["score"]=float(sc); res.append(hit)
    return res

def build_context(hits): return "\n\n".join([f"[{h['id']}] {h['text']}" for h in hits])

# ---------------- SYSTEM PROMPT ----------------
SYSTEM_PROMPT = """
The Value Chain Builder is a seasoned management consultant specializing in global markets, value chains, and supply chain optimization, delivering data-driven analyses aligned with uploaded reference documents. It assesses cost structures and industry interdependencies globally, regionally, and locally.

**Analytical Framework**:

1. **Global Overview**:
   1.1 **Introduction**:
   Detailed overview of [target industry]:

* **Market Size & Growth**: Current size, past decade growth, projections (5–10 years).
* **Key Milestones**: 8–10 historical milestones (innovations, policies, disruptions) in a table.
* **Industry Segments & Shares**: Segment characteristics, global market shares.
* **Emerging Trends**: 10–12 quantified emerging trends categorized in a table.

1.2 **Global Trade**:
Analysis of trade dynamics:

* **Leading Countries**: Top 5 countries (2023 import/export values, production tonnes/year).
* **Major Insights**: 3 detailed insights per top 5 countries (growth factors, technologies, differentiators).
* **Top Global Suppliers**: 7 key suppliers, market shares, capacities, competitive advantages.
* **Case Study (2 Countries)**: Pause & confirm before proceeding:

  * **GDP & Industry Contribution**: GDP value (USD), industry % of GDP.
  * **Workforce**: Employment & workforce %.
  * **Market Size**: Production, exports, turnover (USD), OEMs/component manufacturers presence.
  * **Top Players**: 3–5 OEMs, 8–10 component manufacturers.
  * **Key Milestones**: 8–10 detailed milestones in a table.

2. **Country Landscape Overview**:
   Analyze [target industry] in user-specified country:

* **Local Segmentation**: 4–6 manufacturers, units sold (2023), imports (local currency).
* **Key Players**: Top 6–8 suppliers, distributors, manufacturers.
* **Market Shares**: Market shares by segment.
* **Demand Projections**: Demand for 2024, 2030, 2035, 2040.

2.2 **Segments Overview**:
Major segments breakdown:

* **Import Dependency**: Segment imports (local currency), % total imports.
* **Market Size**: Growth rates, segment shares.
* **Trends & Insights**: Import trends (2018–2023), 5 demand drivers.

2.3 **Trade Analysis**:
Evaluate country’s industry trade:

* **Top Exporting Countries**: Leading exporters (2018–2023) by segment.
* **Rankings**: Top 5 exporting countries.
* **Trends**: 3 key trade trends (5–10 years).

2.4 **Local Capabilities**:
Production & supply chain:

* **Current Capabilities**: Local players, ongoing projects (production, partnerships, JVs, takeovers).
* **Upcoming Suppliers**: Emerging suppliers categories.

3. **Associated Industry & Value Chain Analysis**:
   3.1 **Associated Industry Analysis**:
   Select 4 industries from provided "THE INDUSTRY LIST" (platform & midstream):

* **Industry List**: 4 selected industries table.
* **Industry Role**: Roles & value chain contributions.
* **Supplier Tiers**: Top 4 suppliers per industry (2 local, 2 global), specific company names only.
* **Cost Contribution**: Cost contribution estimates & 3 insights per industry.

3.2 **Value Chain Analysis**:
Prioritize 2 industries with highest costs:

* **4-Step Value Chain**: Raw materials, processing, component production, assembly:

  * 3 technology needs
  * 3 trends
  * Local landscape
  * 3 barriers
  * 3 potential strategic advantages
* **Supplier Landscape**: Top 3 global & local suppliers per stage.
* **Supplier Case Study**: Pause & confirm, then provide:

  * Name, HQ, founding year
  * Revenue, top 3 global locations
  * Employee count, top 3 products/services, value contributions.
* **Localization Opportunities**: 4 opportunities aligning with national industrial goals.

3.3 **Raw Material Analysis**:
Select 2 raw material industries (Upstream preferred) from "THE INDUSTRY LIST":

* **4-Step Value Chain**: Analyze each stage (raw materials, processing, component, assembly):

  * 3 technology needs
  * 3 trends
  * Local landscape
  * 3 barriers
  * 3 potential strategic advantages
* **Supplier Landscape**: Top 3 global & local suppliers per stage.
* **Raw Material Supplier Case Study**: Pause & confirm, then provide detailed:

  * Name, HQ, founding year
  * Revenue, top 3–5 operating locations
  * Employee count, top 3 products/services, value contributions.
* **Localization Opportunities**: 4 localization opportunities aligned with national goals.

4. **Consolidated Opportunities**:
   Summarize actionable recommendations for investments, localization, and partnerships aligned with industry trends & regional priorities (provide thorough examples).

Before starting, confirm:

* Target industry and specified country.
* Layout and points to cover (State the Titles of the categories and subcategories as the layout of the analysis then proceed to do each section)
* After each main section, confirm thoroughness. Pause & seek confirmation for detailed, lengthy subsections. Populate accurate numbers & mirror uploaded reference files' analytical depth and narrative style.

Note that if the user didn’t give an input for supplier’s case studies, choose random top suppliers.
Note that you need to always return case studies for two suppliers, if the user gave one only, chose a random top supplier as the second one.
Note that you should insure that you alway spopulate numbers, make sure NOT to put "XX", ALWAYS populate the figures.
All analyses must **not default to KSA**. Always ask the user to confirm the **target country** and **target industry** before proceeding with any section. Do not assume a specific geography unless explicitly instructed by the user.
"""

SECTIONS = [
    ("1.1 Introduction", "Produce Section 1.1 (Introduction). Include: market size & growth; key milestones table (8–10); segments & shares; emerging trends table (10–12). Keep concise; cite sources like [file#chunk]."),
    ("1.2 Global Trade", "Produce Section 1.2 (Global Trade). Include: leading countries (2023 import/export values, production); 3 insights per top 5; top global suppliers (7) with shares/capacities. END by asking for confirmation for 'Case Study (2 Countries)'."),
    ("1.2 Case Study (2 Countries)", "ONLY IF CONFIRMED. Provide 2-country case study with GDP & industry contribution; workforce; market size; top players; milestones table. Keep under cap; cite clearly."),
    ("2.1 Country Landscape", "Produce 2.1 Country Landscape: local segmentation; key players; shares; demand projections (2024, 2030, 2035, 2040)."),
    ("2.2 Segments Overview", "Produce 2.2 Segments Overview: import dependency; segment market size & growth; import trends (2018–2023); 5 demand drivers."),
    ("2.3 Trade Analysis", "Produce 2.3: top exporting countries by segment (2018–2023); top 5 exporters; 3 key trade trends."),
    ("2.4 Local Capabilities", "Produce 2.4: current capabilities; ongoing projects (production, partnerships, JVs, takeovers); upcoming suppliers."),
    ("3.1 Associated Industries", "Select 4 industries from THE INDUSTRY LIST (platform & midstream). Roles; supplier tiers (2 local, 2 global each); cost contribution + 3 insights each."),
    ("3.2 Value Chain Analysis", "Pick the 2 highest-cost industries. Build 4-step value chain; technology needs (3); trends (3); local landscape; barriers (3); strategic advantages (3). Supplier landscape (3 global + 3 local per stage). Ask to confirm 'Supplier Case Study'."),
    ("3.2 Supplier Case Study", "ONLY IF CONFIRMED. Detailed supplier case study: name, HQ, founding year, revenue, top 3 locations, employees, top 3 products/services, value contributions. ALWAYS return two suppliers; if user gave one, pick a top second."),
    ("3.3 Raw Material Analysis", "Select 2 upstream raw material industries. Same 4-step analysis; supplier landscape. Ask to confirm 'Raw Material Supplier Case Study'."),
    ("3.3 Raw Material Supplier Case Study", "ONLY IF CONFIRMED. Detailed raw material supplier case study: name, HQ, founding year, revenue, 3–5 locations, employees, top 3 products/services, value contributions."),
    ("4. Consolidated Opportunities", "Summarize actionable opportunities for investments/localization/partnerships aligned with national goals. Include concrete examples. Keep crisp."),
]

# ---------- Section focus validator ----------
FORBIDDEN_BY_SECTION = {
    "3.1 Associated Industries": [
        "layout of the analysis", "1. global overview", "1.1 introduction", "1.2 global trade"
    ],
    "3.2 Value Chain Analysis": ["layout of the analysis", "global overview", "1.1 introduction"],
    "3.2 Supplier Case Study": ["layout of the analysis", "global overview"],
    "3.3 Raw Material Analysis": ["layout of the analysis", "global overview"],
    "3.3 Raw Material Supplier Case Study": ["layout of the analysis", "global overview"],
}
REQUIRED_HINTS = {
    "3.1 Associated Industries": ["associated", "supplier", "industry"],
}

def violates_section_focus(name: str, text: str) -> bool:
    t = text.lower()
    for bad in FORBIDDEN_BY_SECTION.get(name, []):
        if bad in t:
            return True
    reqs = REQUIRED_HINTS.get(name)
    if reqs and not any(r in t for r in reqs):
        return True
    return False

# ---------- Tag helpers & mappings ----------
def normalize_tag(tag: Optional[str]) -> str:
    if not tag: return ""
    t = tag.upper()
    t = re.sub(r'[^A-Z0-9]+', '_', t).strip('_')
    return t

PRECASE_TO_TAG = { 1: "CASE_STUDY_2_COUNTRIES", 8: "SUPPLIER_CASE_STUDY", 10: "RAW_MATERIAL_CASE_STUDY" }
CASE_IDX_TO_TAG = { 2: "CASE_STUDY_2_COUNTRIES", 9: "SUPPLIER_CASE_STUDY", 11: "RAW_MATERIAL_CASE_STUDY" }
CONFIRM_PHRASE_RE = re.compile(r'\b(please confirm|confirm if|would you like to proceed)\b', re.I)
PLACEHOLDER_RE = re.compile(r'\[\s*Name\b|\[HQ\b|\[Revenue\b|\bHere[’\'`s ]+the outline\b', re.I)

def gate_to_index(tag: str) -> int:
    t = normalize_tag(tag)
    if t.startswith("CASE_STUDY"): return 2
    if t.startswith("SUPPLIER_CASE_STUDY"): return 9
    if t.startswith("RAW_MATERIAL_CASE_STUDY"): return 11
    return st.session_state.current_idx

def confirm_text_for_tag(tag: str, specifics: Optional[str] = None) -> str:
    t = normalize_tag(tag)
    extra = f" User specifics: {specifics.strip()}" if specifics and specifics.strip() else ""
    if t.startswith("SUPPLIER_CASE_STUDY"):
        return ("CONFIRM: Proceed with 3.2 Supplier Case Study. "
                "Always return TWO supplier case studies. If only one supplier is provided, pick a second top supplier automatically."
                + extra)
    elif t.startswith("RAW_MATERIAL_CASE_STUDY"):
        return ("CONFIRM: Proceed with 3.3 Raw Material Supplier Case Study. "
                "If no raw material/supplier is provided, choose a notable representative example."
                + extra)
    elif t.startswith("CASE_STUDY"):
        return ("CONFIRM: Proceed with 1.2 Case Study (2 Countries). "
                "If countries are not specified, select two most relevant based on trade/production/suppliers."
                + extra)
    else:
        return "CONFIRM: Proceed with the pending section." + extra

# ---------------- LLM with auto-continue ----------------
def call_llm(section_name, section_instruction, user_notes, target_country, target_industry, k=RETRIEVAL_K, temperature=0.1):
    index, docs = load_index_cached()
    hits = retrieve(index, docs, f"{target_industry} in {target_country} :: {section_instruction} :: {user_notes}", k=k)
    context = build_context(hits)

    formatting_rules = (
        "FORMAT RULES:\n"
        "- Never output placeholders like XX, TBD, —; provide a numeric estimate with units and mark '(estimate)' if uncertain.\n"
        "- Do NOT use raw HTML in tables (e.g., <br>, <p>). Inside table cells, separate items with '; '.\n"
    )
    if section_name == "1.1 Introduction":
        section_guard = "HARD CONSTRAINT: You may include the analysis layout here only if needed."
    else:
        section_guard = ("HARD CONSTRAINTS:\n"
                         f"- Output ONLY the content for '{section_name}'.\n"
                         "- Do NOT include any 'Layout of the Analysis', 'Global Overview', or other sections.\n")

    base_messages = [
        {"role":"system","content": SYSTEM_PROMPT},
        {"role":"user","content": (
            f"TARGET INDUSTRY: {target_industry}\nTARGET COUNTRY: {target_country}\n\n"
            f"RETRIEVED EXCERPTS (use and cite):\n{context}\n\n"
            f"SECTION TO PRODUCE (only this):\n{section_instruction}\n\n"
            f"{section_guard}\n{formatting_rules}"
            f"USER NOTES/QUESTION:\n{user_notes}\n"
        )}
    ]

    budget = SECTION_BUDGET.get(section_name, 1440)
    out_parts=[]; total={"prompt":0,"completion":0,"total":0}

    resp = client.chat.completions.create(model=CHAT_MODEL, messages=base_messages, temperature=temperature, max_tokens=budget)
    text = resp.choices[0].message.content or ""
    out_parts.append(text)
    fin = getattr(resp.choices[0], "finish_reason", None)
    u = getattr(resp, "usage", None)
    if u:
        total["prompt"]+=getattr(u,"prompt_tokens",0)
        total["completion"]+=getattr(u,"completion_tokens",0)
        total["total"]+=getattr(u,"total_tokens",0)

    rounds=0
    while fin == "length" and rounds < MAX_CONT_ROUNDS:
        rounds += 1
        cont = [
            {"role":"system","content": SYSTEM_PROMPT},
            {"role":"user","content":
             "CONTINUE exactly where you stopped. Do NOT repeat earlier content. "
             "If a table is mid-way, continue rows under same headers. Output only continuation."},
            {"role":"assistant","content": text[-800:]}
        ]
        resp = client.chat.completions.create(model=CHAT_MODEL, messages=cont, temperature=temperature, max_tokens=int(budget*0.8))
        text = resp.choices[0].message.content or ""
        out_parts.append(text)
        fin = getattr(resp.choices[0], "finish_reason", None)
        u = getattr(resp, "usage", None)
        if u:
            total["prompt"]+=getattr(u,"prompt_tokens",0)
            total["completion"]+=getattr(u,"completion_tokens",0)
            total["total"]+=getattr(u,"total_tokens",0)

    final = "\n".join(out_parts)
    st.session_state.setdefault("_logs", []).append(
        {"ts": time.time(), "section": section_name, **total,
         "top_sources":[h["id"] for h in hits[:6]]}
    )
    return final

# ---------- Presentation cleanup ----------
READY_LINE_RE = re.compile(r'^.*<READY:[^>]+>\??.*\n?', flags=re.MULTILINE)

def sanitize_markdown(s: str) -> str:
    s = READY_LINE_RE.sub('', s)                                  # drop READY lines
    s = re.sub(r'(?is)<\s*br\s*/?\s*>', ' ; ', s)                  # <br> variants → '; '
    s = s.replace('&lt;br&gt;', ' ; ').replace('&lt;BR&gt;', ' ; ') # encoded br
    s = re.sub(r'(?is)</?\s*p[^>]*>', ' ', s)                      # <p> tags
    s = re.sub(r'[ \t]{2,}', ' ', s)                               # collapse spaces
    return s

# ---------- Section add with de-dupe ----------
def add_section(name: str, raw_text: str):
    text = sanitize_markdown(raw_text)
    if st.session_state.sections_done and st.session_state.sections_done[-1][0] == name:
        st.session_state.sections_done[-1] = (name, text)
    else:
        st.session_state.sections_done.append((name, text))

# ---------------- Outline (simple & clean) ----------------
def render_outline():
    titles = [name for name,_ in SECTIONS]
    with st.expander("Analysis Structure", expanded=False):
        st.markdown("- " + "\n- ".join(titles))

# ---------------- Export to Word (appears only at the very end) ----------------
def sections_to_docx_bytes(sections: List[Tuple[str,str]]) -> bytes:
    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
    for name, content in sections:
        h = doc.add_heading(level=1); h_run = h.add_run(name); h_run.font.color.rgb = None
        for block in content.split("\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("|") and block.count("|") >= 2:
                p = doc.add_paragraph()
                p.add_run(block)
            else:
                doc.add_paragraph(block)
        doc.add_paragraph()
    bio = BytesIO(); doc.save(bio); bio.seek(0)
    return bio.getvalue()

# ---------------- UI ----------------
render_outline()

colA, colB = st.columns([1,1])
with colA: target_industry = st.text_input("Target industry", key="ti")
with colB: target_country = st.text_input("Target country", key="tc")
user_notes = st.text_area("Analyst notes (optional)", key="notes", height=80,
                          placeholder="Focus on post-2020 trends, policy levers, localization levers…")

# Always visible switch-style checkbox
expand_all = st.checkbox("Expand all sections", value=False, key="expand_all")

generate_clicked = st.button("Generate / Proceed", type="primary", key="btn_generate")

# session
ss = st.session_state
ss.setdefault("sections_done", [])
ss.setdefault("current_idx", 0)
ss.setdefault("pending_tag", None)
ss.setdefault("completed_tags", set())
ss.setdefault("_just_confirmed", False)

# ---------- Case-study & section focus runners ----------
def enforce_section_focus(name: str, text: str, notes: str, target_country: str, target_industry: str, instr: str) -> str:
    if not violates_section_focus(name, text):
        return text
    fix_notes = (notes + "\n\n"
                 f"FIX: Your previous output included unrelated content (e.g., layout/global overview). "
                 f"Rewrite strictly for '{name}' only. Do NOT include any other sections. "
                 "Follow the required fields for this section and the formatting rules.")
    repaired = call_llm(name, instr, fix_notes, target_country, target_industry)
    return repaired

def run_idx(idx: int, notes: str):
    autoconf = False
    if idx in CASE_IDX_TO_TAG and not ss._just_confirmed:
        tag = CASE_IDX_TO_TAG[idx]
        if normalize_tag(tag) not in ss.completed_tags:
            notes = notes + "\n\n" + confirm_text_for_tag(tag)
            autoconf = True

    name, instr = SECTIONS[idx]
    with st.spinner(f"Working on {name}…"):
        raw_out = call_llm(name, instr, notes, target_country, target_industry)

    # Force completion for case-studies if outline/confirm/placeholder shows
    if idx in CASE_IDX_TO_TAG and (CONFIRM_PHRASE_RE.search(raw_out) or PLACEHOLDER_RE.search(raw_out)):
        tag = CASE_IDX_TO_TAG[idx]
        force_directive = (
            f"CONFIRMED: Generate the full {name} now. "
            "Do not ask for confirmation. Produce the complete case study as requested:\n"
        )
        if "Supplier" in name:
            force_directive += (
                "- Return TWO suppliers. If the user provided only one or none, select top suppliers relevant to the target country/industry.\n"
                "- For each supplier include: Name; HQ; founding year; revenue (USD, latest); top 3 global locations; employee count; top 3 products/services; value contributions.\n"
            )
        elif "Raw Material" in name:
            force_directive += (
                "- Return TWO raw material suppliers (or representative leaders) relevant to the target country/industry.\n"
                "- For each: Name; HQ; founding year; revenue (USD); top 3–5 operating locations; employee count; top 3 products/services; value contributions.\n"
            )
        else:  # 1.2 for two countries
            force_directive += (
                "- Provide two countries with: GDP & industry contribution; workforce & %; market size/production/exports/turnover; 3–5 OEMs & 8–10 components players; milestones table.\n"
            )
        force_directive += "Use the same formatting rules. No placeholders like XX; no HTML tags. Output the final content."
        raw_out = call_llm(name, force_directive, notes, target_country, target_industry)
        ss.completed_tags.add(normalize_tag(tag))
        ss.pending_tag = None
        ss._just_confirmed = False

    raw_out = enforce_section_focus(name, raw_out, notes, target_country, target_industry, instr)

    asked_confirm = bool(CONFIRM_PHRASE_RE.search(raw_out))
    if idx in CASE_IDX_TO_TAG and (ss._just_confirmed or autoconf):
        asked_confirm = False
        ss.completed_tags.add(normalize_tag(CASE_IDX_TO_TAG[idx]))

    if asked_confirm and idx in CASE_IDX_TO_TAG:
        tag = CASE_IDX_TO_TAG[idx]
        if normalize_tag(tag) not in ss.completed_tags:
            ss.pending_tag = tag
        else:
            ss.pending_tag = None
    else:
        tags = re.findall(r"<READY:\s*([A-Za-z0-9_ \-]+)\s*\?>", raw_out)
        next_pending = None
        for raw in tags[::-1]:
            t = normalize_tag(raw)
            if t in {"SUPPLIER_CASE_STUDY", "RAW_MATERIAL_CASE_STUDY", "CASE_STUDY_2_COUNTRIES"}:
                if t not in ss.completed_tags:
                    next_pending = t
                    break
        ss.pending_tag = next_pending

    add_section(name, raw_out)
    ss._just_confirmed = False

# ensure index buildable
try: load_index_cached()
except Exception as e: st.warning(str(e))

# main driver button
if generate_clicked:
    if not target_industry or not target_country:
        st.warning("Please provide both target industry and target country.")
    else:
        if ss.pending_tag:
            gi = gate_to_index(ss.pending_tag)
            confirm_blob = confirm_text_for_tag(ss.pending_tag)
            ss._just_confirmed = True
            run_idx(gi, user_notes + "\n\n" + confirm_blob)
            ss.completed_tags.add(normalize_tag(ss.pending_tag))
            ss.pending_tag = None
            ss.current_idx = max(ss.current_idx, gi + 1)
            st.rerun()
        else:
            if ss.current_idx < len(SECTIONS):
                run_idx(ss.current_idx, user_notes)
                ss.current_idx += 1
                st.rerun()

# render sections + gating under the LAST one only
for i,(name,text) in enumerate(ss.sections_done):
    last = i == len(ss.sections_done)-1
    with st.expander(name, expanded=expand_all or last):
        st.markdown(text)

        sec_idx = next((idx for idx,(n,_) in enumerate(SECTIONS) if n == name), None)

        implicit_tag = None
        if ss.get("pending_tag") is None and sec_idx in PRECASE_TO_TAG:
            candidate = PRECASE_TO_TAG[sec_idx]
            if normalize_tag(candidate) not in ss.completed_tags:
                implicit_tag = candidate
        if ss.get("pending_tag") is None and sec_idx in CASE_IDX_TO_TAG and CONFIRM_PHRASE_RE.search(text):
            candidate = CASE_IDX_TO_TAG[sec_idx]
            if normalize_tag(candidate) not in ss.completed_tags:
                implicit_tag = candidate

        effective_tag = ss.get("pending_tag") or implicit_tag

        if last and effective_tag:
            norm_tag = normalize_tag(effective_tag)
            st.info(f"{norm_tag.replace('_',' ').title()} — provide details (optional) or proceed with defaults.")
            gate_msg = st.text_input("Your input (optional)", key=f"gate_input_{i}_{norm_tag}")
            c1, c2 = st.columns(2)
            with c1:
                if st.button("Proceed with input", key=f"btn_gate_with_{i}_{norm_tag}"):
                    gi = gate_to_index(effective_tag)
                    confirm_blob = confirm_text_for_tag(effective_tag, specifics=gate_msg)
                    ss._just_confirmed = True
                    run_idx(gi, user_notes + (("\n\nUSER SPECIFICS:\n"+gate_msg.strip()) if gate_msg.strip() else "") + "\n\n" + confirm_blob)
                    ss.completed_tags.add(norm_tag)
                    ss.pending_tag = None
                    ss.current_idx = max(ss.current_idx, gi + 1)
                    st.rerun()
            with c2:
                if st.button("Proceed (default)", key=f"btn_gate_default_{i}_{norm_tag}"):
                    gi = gate_to_index(effective_tag)
                    ss._just_confirmed = True
                    run_idx(gi, user_notes + "\n\n" + confirm_text_for_tag(effective_tag))
                    ss.completed_tags.add(norm_tag)
                    ss.pending_tag = None
                    ss.current_idx = max(ss.current_idx, gi + 1)
                    st.rerun()
        elif last:
            if ss.current_idx < len(SECTIONS):
                nxt_name, _ = SECTIONS[ss.current_idx]
                extra_notes = st.text_input("Optional note to apply to the next section", key=f"next_note_{i}")
                if st.button(f"Proceed: {nxt_name}", key=f"btn_next_{i}_{ss.current_idx}"):
                    run_idx(ss.current_idx, user_notes + (("\n\nEXTRA NOTE:\n"+extra_notes.strip()) if extra_notes.strip() else ""))
                    ss.current_idx += 1
                    st.rerun()

        # Export only after the final section
        if last and name == "4. Consolidated Opportunities" and ss.current_idx >= len(SECTIONS):
            docx_bytes = sections_to_docx_bytes(ss.sections_done)
            st.download_button(
                "Export full analysis (Word)",
                data=docx_bytes,
                file_name=f"Efficio_Value_Chain_Analysis_{time.strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                key="btn_export_docx_final"
            )

# footer (Efficio-aligned disclaimer)
st.markdown(
    "<div class='ef-hint'>The Value Chain Builder is built on Efficio’s proven value chain methodology. "
    "Outputs are AI-generated from available sources and may contain inaccuracies—please review and validate key figures.</div>",
    unsafe_allow_html=True
)
